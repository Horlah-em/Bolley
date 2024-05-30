from pydrive2.auth import GoogleAuth
from pydrive2.drive import GoogleDrive
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import Pinecone as LangchainPinecone
from langchain.docstore.document import Document  # Import the correct Document class
from dotenv import load_dotenv
import os
from pinecone import Pinecone, ServerlessSpec
import time
import urllib3.exceptions
from langchain.embeddings import HuggingFaceEmbeddings

# Custom DocxLoader class
class DocxLoader:
    def __init__(self, file_path):
        self.file_path = file_path

    def load(self):
        from docx import Document as DocxDocument
        doc = DocxDocument(self.file_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        metadata = {'file_path': self.file_path}
        return [Document(page_content='\n'.join(text), metadata=metadata)]

# Step 1: Authenticate and Initialize PyDrive
def authenticate_and_initialize_drive():
    gauth = GoogleAuth()
    gauth.LocalWebserverAuth()  # Creates local webserver and automatically handles authentication.
    drive = GoogleDrive(gauth)
    return drive

# Step 2: Download files from Google Drive
def download_files_from_drive(drive, folder_id):
    file_list = drive.ListFile({'q': f"'{folder_id}' in parents and trashed=false"}).GetList()
    downloaded_files = []
    for file in file_list:
        file_id = file['id']
        file_name = file['title']
        file.GetContentFile(file_name)
        downloaded_files.append(file_name)
    return downloaded_files

# Step 3: Load documents from downloaded files
def load_docs_from_files(file_paths):
    documents = []
    for file_path in file_paths:
        if file_path.endswith('.docx'):
            loader = DocxLoader(file_path)
        else:
            continue  # Skip unsupported file types
        documents.extend(loader.load())
    return documents

# Step 4: Split documents
def split_docs(documents, chunk_size=500, chunk_overlap=20):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    docs = text_splitter.split_documents(documents)
    return docs

# Retry function to handle transient errors
def retry_function(func, max_retries=5, backoff_factor=0.5):
    for attempt in range(max_retries):
        try:
            return func()
        except (urllib3.exceptions.ProtocolError, ConnectionResetError) as e:
            print(f"Attempt {attempt + 1} failed with error: {e}. Retrying...")
            time.sleep(backoff_factor * (2 ** attempt))
    raise Exception(f"Function failed after {max_retries} retries")

# Step 5: Main function to execute the workflow
def main():
    drive = authenticate_and_initialize_drive()
    folder_id = '1f1oAM687MypemrSGhM4dGZhUEJpfuQfX'  # Replace with your folder ID
    downloaded_files = download_files_from_drive(drive, folder_id)
    documents = load_docs_from_files(downloaded_files)
    print(f"Loaded {len(documents)} documents.")
    
    docs = split_docs(documents)
    print(f"Split into {len(docs)} chunks.")
    
    # Load environment variables
    load_dotenv()
    OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
    PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
    PINECONE_ENV = os.getenv("PINECONE_ENV")
    
    # Initialize Pinecone
    pc = Pinecone(api_key=PINECONE_API_KEY)
    
    # Create the index if it doesn't exist
    index_name = "medical-ai"
    if index_name in pc.list_indexes().names():
        pc.delete_index(index_name)
    
    pc.create_index(
        name=index_name,
        dimension=1536,  # Correct dimension for the embedding model
        metric='euclidean',
        spec=ServerlessSpec(
            cloud='aws',
            region='us-east-1'
        )
    )
    
    # Create embeddings
    embeddings = OpenAIEmbeddings(model="text-embedding-ada-002", openai_api_key=OPENAI_API_KEY)
    
    # Function to add texts to Pinecone with retry logic
    def add_texts_to_pinecone():
        return LangchainPinecone.from_documents(docs, embeddings, index_name=index_name)
    
    # Index documents using Langchain's Pinecone wrapper with retry
    langchain_index = retry_function(add_texts_to_pinecone)
    
    # Example query function
    def get_similar_docs(query, k=1, score=False):
        if score:
            similar_docs = langchain_index.similarity_search_with_score(query, k=k)
        else:
            similar_docs = langchain_index.similarity_search(query, k=k)
        return similar_docs
    
    # Example usage
    query = "Your query here"
    results = get_similar_docs(query, k=3, score=True)
    for result in results:
        print(result)
    
if __name__ == "__main__":
    main()
